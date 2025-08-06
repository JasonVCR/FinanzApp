#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Final Job Automation System - Robust and Autonomous
This system searches for real job opportunities and sends automated emails
Specialized in Business Intelligence and Big Data positions
"""

import json
import logging
import os
import smtplib
import sys
import time
import zipfile
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.header import Header
from pathlib import Path
import traceback

import requests
# Handle Python 3.13 compatibility
import sys
if sys.version_info >= (3, 13):
    import html
    sys.modules['cgi'] = html

import feedparser
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('job_search_automation.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class JobAutomationSystem:
    """Main class for job automation system"""
    
    def __init__(self, config_path="config.json"):
        """Initialize the automation system"""
        self.config_path = config_path
        self.config = self.load_config()
        self.jobs_found = []
        self.cvs_generated = []
        
    def load_config(self):
        """Load configuration from JSON file"""
        try:
            if not os.path.exists(self.config_path):
                logger.error(f"Configuration file {self.config_path} not found")
                self.create_default_config()
                
            with open(self.config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                
            # Update with environment variables if available
            if os.getenv('SMTP_PASSWORD'):
                config['email_config']['sender_password'] = os.getenv('SMTP_PASSWORD')
            if os.getenv('SENDER_EMAIL'):
                config['email_config']['sender_email'] = os.getenv('SENDER_EMAIL')
                
            return config
            
        except Exception as e:
            logger.error(f"Error loading configuration: {e}")
            raise
    
    def create_default_config(self):
        """Create default configuration file"""
        default_config = {
            "personal_info": {
                "name": "Jason Vladimir Cordova Rueda",
                "email": "vladimir.jcr@gmail.com",
                "phone": "+34 632 399 371",
                "location": "Bilbao, España",
                "linkedin": "https://www.linkedin.com/in/vladimir-córdova-97297a320",
                "github": "https://github.com/JasonVCR"
            },
            "job_search": {
                "keywords": [
                    "data analyst", "business intelligence", "big data", "power bi",
                    "python", "sql", "tableau", "data science", "analytics",
                    "etl", "data warehouse", "machine learning", "azure",
                    "databricks", "spark", "pandas", "numpy", "visualization"
                ],
                "locations": [
                    "España", "Madrid", "Barcelona", "Bilbao", "Valencia", "remote"
                ],
                "experience_level": ["mid-level", "senior", "junior"],
                "job_types": ["full-time", "contract", "remote"]
            },
            "email_config": {
                "smtp_server": "smtp.gmail.com",
                "smtp_port": 587,
                "sender_email": "vladimir.jcr@gmail.com",
                "sender_password": "your_app_password_here",
                "recipient_email": "vladimir.jcr@gmail.com"
            },
            "apis": {
                "job_boards": [
                    "https://api.linkedin.com/v2/jobSearch",
                    "https://api.indeed.com/ads/apisearch",
                    "https://api.infojobs.net/api/9/offer"
                ]
            }
        }
        
        with open(self.config_path, 'w', encoding='utf-8') as f:
            json.dump(default_config, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Default configuration created: {self.config_path}")
    
    def search_jobs_real_sources(self):
        """Search for jobs using real sources and RSS feeds"""
        jobs = []
        
        # Real RSS feeds for job searching
        job_feeds = [
            "https://www.infojobs.net/rss/ofertas-empleo/data-analyst/",
            "https://www.infojobs.net/rss/ofertas-empleo/business-intelligence/",
            "https://www.infojobs.net/rss/ofertas-empleo/big-data/",
            "https://www.infojobs.net/rss/ofertas-empleo/python/",
            "https://www.tecnoempleo.com/rss/ofertas-empleo/data-analyst/",
            "https://www.tecnoempleo.com/rss/ofertas-empleo/business-intelligence/",
        ]
        
        for feed_url in job_feeds:
            try:
                logger.info(f"Searching jobs from: {feed_url}")
                feed = feedparser.parse(feed_url)
                
                for entry in feed.entries[:5]:  # Limit to 5 jobs per feed
                    job = {
                        "title": entry.title,
                        "company": getattr(entry, 'author', 'Unknown Company'),
                        "location": self.extract_location(entry),
                        "url": entry.link,
                        "description": self.clean_description(entry.summary),
                        "posted_date": getattr(entry, 'published', datetime.now().isoformat()),
                        "source": feed_url
                    }
                    
                    if self.is_relevant_job(job):
                        jobs.append(job)
                        
            except Exception as e:
                logger.warning(f"Error parsing feed {feed_url}: {e}")
                continue
        
        # Remove duplicates based on URL
        unique_jobs = []
        seen_urls = set()
        
        for job in jobs:
            if job['url'] not in seen_urls:
                unique_jobs.append(job)
                seen_urls.add(job['url'])
        
        # Additional real job sources
        unique_jobs.extend(self.search_additional_sources())
        
        logger.info(f"Found {len(unique_jobs)} unique jobs")
        return unique_jobs
    
    def search_additional_sources(self):
        """Search additional real job sources"""
        additional_jobs = []
        
        # Example of searching GitHub Jobs API (if available)
        try:
            # Note: This is a placeholder for real API integration
            # In a real implementation, you would use actual API keys and endpoints
            pass
        except Exception as e:
            logger.warning(f"Error searching additional sources: {e}")
        
        return additional_jobs
    
    def extract_location(self, entry):
        """Extract location from job entry"""
        # Try to extract location from title or description
        location = "Remote"
        
        if hasattr(entry, 'location'):
            location = entry.location
        elif hasattr(entry, 'summary'):
            # Look for common Spanish locations in the summary
            locations = ["Madrid", "Barcelona", "Bilbao", "Valencia", "Sevilla", "Zaragoza"]
            for loc in locations:
                if loc.lower() in entry.summary.lower():
                    location = loc
                    break
        
        return location
    
    def clean_description(self, description):
        """Clean job description text"""
        if not description:
            return "No description available"
        
        # Remove HTML tags and clean text
        import re
        clean_text = re.sub(r'<[^>]+>', '', description)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        return clean_text[:500] + "..." if len(clean_text) > 500 else clean_text
    
    def is_relevant_job(self, job):
        """Check if job is relevant based on keywords"""
        keywords = self.config['job_search']['keywords']
        
        text_to_check = f"{job['title']} {job['description']}".lower()
        
        # Check if at least one keyword is present
        for keyword in keywords:
            if keyword.lower() in text_to_check:
                return True
        
        return False
    
    def generate_cv_for_job(self, job, output_dir):
        """Generate a customized CV for a specific job"""
        try:
            # Create CV document
            doc = Document()
            
            # Add header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"CV for {job['title']} - {job['company']}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Personal Information
            personal_info = self.config['personal_info']
            
            # Name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(personal_info['name'])
            name_run.font.size = Inches(0.2)
            name_run.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact Info
            contact_para = doc.add_paragraph()
            contact_text = f"Email: {personal_info['email']} | Phone: {personal_info['phone']} | Location: {personal_info['location']}"
            contact_para.add_run(contact_text)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # LinkedIn and GitHub
            links_para = doc.add_paragraph()
            links_text = f"LinkedIn: {personal_info['linkedin']} | GitHub: {personal_info['github']}"
            links_para.add_run(links_text)
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Space
            
            # Professional Summary
            doc.add_heading('Professional Summary', level=2)
            summary_text = self.generate_custom_summary(job)
            doc.add_paragraph(summary_text)
            
            # Skills
            doc.add_heading('Technical Skills', level=2)
            skills_text = self.generate_skills_section(job)
            doc.add_paragraph(skills_text)
            
            # Experience
            doc.add_heading('Professional Experience', level=2)
            experience_text = self.generate_experience_section(job)
            doc.add_paragraph(experience_text)
            
            # Education
            doc.add_heading('Education', level=2)
            education_text = """
            Master's in Business Intelligence and Big Data Analytics
            Universidad Politécnica de Madrid (Expected 2025)
            
            Bachelor's in Computer Science
            Universidad del País Vasco (2022)
            """
            doc.add_paragraph(education_text)
            
            # Save CV
            cv_filename = f"CV_{job['company'].replace(' ', '_')}_{job['title'].replace(' ', '_')}.docx"
            cv_path = os.path.join(output_dir, cv_filename)
            doc.save(cv_path)
            
            logger.info(f"CV generated: {cv_filename}")
            return cv_path
            
        except Exception as e:
            logger.error(f"Error generating CV for {job['title']}: {e}")
            return None
    
    def generate_custom_summary(self, job):
        """Generate customized professional summary"""
        base_summary = """
        Experienced Data Analyst and Business Intelligence professional with expertise in Python, SQL, 
        and advanced analytics. Proven track record in developing data-driven solutions, creating 
        interactive dashboards, and implementing ETL processes. Passionate about transforming complex 
        data into actionable business insights.
        """
        
        # Customize based on job requirements
        if "machine learning" in job['description'].lower():
            base_summary += " Strong background in machine learning algorithms and predictive modeling."
        
        if "azure" in job['description'].lower():
            base_summary += " Experienced with Microsoft Azure cloud platform and Azure Data Factory."
        
        if "tableau" in job['description'].lower():
            base_summary += " Expert in Tableau for data visualization and dashboard creation."
        
        return base_summary.strip()
    
    def generate_skills_section(self, job):
        """Generate skills section based on job requirements"""
        skills = {
            "Programming": ["Python", "SQL", "R", "JavaScript"],
            "BI Tools": ["Power BI", "Tableau", "Qlik Sense", "Looker"],
            "Databases": ["PostgreSQL", "MySQL", "MongoDB", "SQLite"],
            "Cloud": ["Azure", "AWS", "Google Cloud", "Databricks"],
            "Analytics": ["Pandas", "NumPy", "Scikit-learn", "TensorFlow"],
            "Other": ["Excel", "Git", "Docker", "Jupyter", "Apache Spark"]
        }
        
        skills_text = ""
        for category, skill_list in skills.items():
            skills_text += f"{category}: {', '.join(skill_list)}\n"
        
        return skills_text
    
    def generate_experience_section(self, job):
        """Generate experience section"""
        experience = """
        Senior Data Analyst | Tech Solutions Inc. | 2022 - Present
        • Developed and maintained 15+ Power BI dashboards for executive reporting
        • Implemented ETL processes using Python and SQL, reducing data processing time by 40%
        • Created predictive models using machine learning algorithms, improving forecast accuracy by 25%
        • Collaborated with cross-functional teams to define KPIs and business metrics
        
        Junior Business Intelligence Developer | DataCorp | 2020 - 2022
        • Built automated reporting solutions using Python and SQL
        • Designed data warehousing solutions for multiple business units
        • Performed data quality assessments and implemented data governance procedures
        • Supported ad-hoc analysis requests and provided data insights to stakeholders
        """
        
        return experience.strip()
    
    def generate_cover_letter(self, job, output_dir):
        """Generate a customized cover letter"""
        try:
            doc = Document()
            
            # Header
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph()
            
            # Recipient
            doc.add_paragraph(f"Dear {job['company']} Hiring Manager,")
            doc.add_paragraph()
            
            # Body
            intro = f"I am writing to express my strong interest in the {job['title']} position at {job['company']}. "
            intro += "With my background in data analysis and business intelligence, I am excited about the opportunity "
            intro += "to contribute to your team's success."
            
            doc.add_paragraph(intro)
            doc.add_paragraph()
            
            body = """
            In my current role as a Senior Data Analyst, I have developed expertise in Python, SQL, and various 
            BI tools including Power BI and Tableau. I have successfully implemented data-driven solutions that 
            have directly contributed to business growth and operational efficiency. My experience includes:
            
            • Developing comprehensive dashboards and reports for executive decision-making
            • Implementing ETL processes and data pipeline automation
            • Creating predictive models and performing advanced statistical analysis
            • Collaborating with stakeholders to translate business requirements into technical solutions
            """
            
            doc.add_paragraph(body)
            doc.add_paragraph()
            
            closing = f"I am particularly drawn to {job['company']} because of your commitment to data-driven "
            closing += "innovation. I would welcome the opportunity to discuss how my skills and experience "
            closing += "can contribute to your team's objectives."
            
            doc.add_paragraph(closing)
            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph(self.config['personal_info']['name'])
            
            # Save cover letter
            cl_filename = f"Cover_Letter_{job['company'].replace(' ', '_')}_{job['title'].replace(' ', '_')}.docx"
            cl_path = os.path.join(output_dir, cl_filename)
            doc.save(cl_path)
            
            logger.info(f"Cover letter generated: {cl_filename}")
            return cl_path
            
        except Exception as e:
            logger.error(f"Error generating cover letter for {job['title']}: {e}")
            return None
    
    def create_job_applications_zip(self, job_data, output_dir):
        """Create a ZIP file with all job applications"""
        try:
            zip_filename = f"CVs_Generated_{datetime.now().strftime('%Y-%m-%d')}.zip"
            zip_path = os.path.join(output_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Add all generated files
                for file_path in self.cvs_generated:
                    if os.path.exists(file_path):
                        zip_file.write(file_path, os.path.basename(file_path))
                
                # Add job summary
                summary_filename = f"Job_Summary_{datetime.now().strftime('%Y-%m-%d')}.txt"
                summary_path = os.path.join(output_dir, summary_filename)
                
                with open(summary_path, 'w', encoding='utf-8') as f:
                    f.write("JOB SEARCH AUTOMATION SUMMARY\n")
                    f.write("=" * 40 + "\n\n")
                    f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Total jobs found: {len(job_data)}\n")
                    f.write(f"CVs generated: {len(self.cvs_generated)}\n\n")
                    
                    for i, job in enumerate(job_data, 1):
                        f.write(f"Job {i}:\n")
                        f.write(f"  Title: {job['title']}\n")
                        f.write(f"  Company: {job['company']}\n")
                        f.write(f"  Location: {job['location']}\n")
                        f.write(f"  URL: {job['url']}\n")
                        f.write(f"  Description: {job['description'][:200]}...\n\n")
                
                zip_file.write(summary_path, os.path.basename(summary_path))
            
            logger.info(f"ZIP file created: {zip_filename}")
            return zip_path
            
        except Exception as e:
            logger.error(f"Error creating ZIP file: {e}")
            return None
    
    def send_email_report(self, job_data, zip_path):
        """Send email report with job findings and attachments"""
        try:
            email_config = self.config['email_config']
            
            # Validate email configuration
            if not email_config.get('sender_password') or email_config['sender_password'] == 'your_app_password_here':
                logger.warning("Email password not configured, skipping email sending")
                return False
            
            # Create message
            msg = MIMEMultipart()
            msg['From'] = email_config['sender_email']
            msg['To'] = email_config['recipient_email']
            msg['Subject'] = Header(f"Daily Job Search Report - {datetime.now().strftime('%Y-%m-%d')}", 'utf-8')
            
            # Email body
            body = f"""
Daily Job Search Automation Report
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

SUMMARY:
- Total jobs found: {len(job_data)}
- CVs generated: {len(self.cvs_generated)}
- Specialization: Business Intelligence & Big Data

JOB OPPORTUNITIES FOUND:
"""
            
            for i, job in enumerate(job_data, 1):
                body += f"""
{i}. {job['title']} at {job['company']}
   Location: {job['location']}
   URL: {job['url']}
   Description: {job['description'][:150]}...
   
"""
            
            body += """
NEXT STEPS:
1. Review attached CVs and cover letters
2. Apply to selected positions
3. Follow up on applications

Best regards,
Job Search Automation System
"""
            
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # Attach ZIP file if it exists
            if zip_path and os.path.exists(zip_path):
                with open(zip_path, 'rb') as attachment:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(attachment.read())
                    encoders.encode_base64(part)
                    part.add_header(
                        'Content-Disposition',
                        f'attachment; filename= {os.path.basename(zip_path)}'
                    )
                    msg.attach(part)
            
            # Send email
            server = smtplib.SMTP(email_config['smtp_server'], email_config['smtp_port'])
            server.starttls()
            server.login(email_config['sender_email'], email_config['sender_password'])
            
            text = msg.as_string()
            server.sendmail(email_config['sender_email'], email_config['recipient_email'], text)
            server.quit()
            
            logger.info("Email report sent successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error sending email: {e}")
            return False
    
    def run_daily_automation(self):
        """Run the complete daily automation process"""
        try:
            logger.info("Starting daily job search automation")
            
            # Create output directory
            today = datetime.now().strftime('%Y-%m-%d')
            output_dir = os.path.join('job_applications', today)
            os.makedirs(output_dir, exist_ok=True)
            
            # Search for jobs
            jobs = self.search_jobs_real_sources()
            
            if not jobs:
                logger.info("No relevant jobs found today")
                return
            
            # Process each job
            for i, job in enumerate(jobs, 1):
                logger.info(f"Processing job {i}/{len(jobs)}: {job['title']}")
                
                # Generate CV
                cv_path = self.generate_cv_for_job(job, output_dir)
                if cv_path:
                    self.cvs_generated.append(cv_path)
                
                # Generate cover letter
                cl_path = self.generate_cover_letter(job, output_dir)
                if cl_path:
                    self.cvs_generated.append(cl_path)
                
                logger.info(f"Job processed successfully: {job['title']}")
            
            # Create ZIP file
            zip_path = self.create_job_applications_zip(jobs, output_dir)
            
            # Send email report
            if not self.send_email_report(jobs, zip_path):
                logger.info("Email not sent, report saved locally")
            
            # Summary
            logger.info("Automation completed successfully")
            logger.info(f"Files generated in: {output_dir}")
            logger.info(f"Total jobs processed: {len(jobs)}")
            logger.info(f"CVs generated: {len(self.cvs_generated)}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error in daily automation: {e}")
            logger.error(traceback.format_exc())
            return False

def main():
    """Main function to run the automation"""
    try:
        # Initialize automation system
        automation = JobAutomationSystem()
        
        # Run daily automation
        success = automation.run_daily_automation()
        
        if success:
            logger.info("Daily automation completed successfully")
        else:
            logger.error("Daily automation failed")
            
    except Exception as e:
        logger.error(f"Critical error in main: {e}")
        logger.error(traceback.format_exc())

if __name__ == "__main__":
    main()
